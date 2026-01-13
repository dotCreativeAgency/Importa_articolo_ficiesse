#!/usr/bin/env python3
"""
Terminale interattivo per esplorare il database degli articoli.
Permette di navigare, cercare e esportare articoli in formato DOCX.

Uso:
    python esplora_articoli.py [database_sqlite]

Esempio:
    python esplora_articoli.py articoli.db
"""

import sqlite3
import sys
import os
import argparse
import logging


def setup_logging(verbose=False, no_emoji=False):
    level = logging.DEBUG if verbose else logging.INFO
    fmt = "%(message)s"
    logging.basicConfig(level=level, format=fmt)


try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from htmldocx import HtmlToDocx

    HTMLDOCX_AVAILABLE = True
except ImportError:
    HTMLDOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup

    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False


class ArticoliExplorer:
    """Classe per esplorare il database degli articoli"""

    def __init__(self, db_path):
        self.db_path = db_path
        self.conn = None
        self.cursor = None
        self.page_size = 10
        self.current_page = 0
        self.current_filter = None
        self.current_search = None

    def connect(self):
        """Connette al database"""
        if not os.path.exists(self.db_path):
            print(f"Errore: Database '{self.db_path}' non trovato!")
            return False
        self.conn = sqlite3.connect(self.db_path)
        self.conn.row_factory = sqlite3.Row
        self.cursor = self.conn.cursor()
        return True

    def close(self):
        """Chiude la connessione"""
        if self.conn:
            self.conn.close()

    def get_total_count(self):
        """Ritorna il numero totale di articoli (con filtri applicati)"""
        query = "SELECT COUNT(*) FROM t_articoli"
        params = []

        if self.current_filter:
            query += " WHERE argomento = ?"
            params.append(self.current_filter)
        elif self.current_search:
            query += " WHERE titolo_articolo LIKE ? OR sotto_titolo LIKE ?"
            params.extend([f"%{self.current_search}%", f"%{self.current_search}%"])

        self.cursor.execute(query, params)
        return self.cursor.fetchone()[0]

    def get_articles_page(self):
        """Ritorna una pagina di articoli"""
        query = (
            "SELECT id_articolo, data, argomento, "
            "titolo_articolo, sotto_titolo, esportato "
            "FROM t_articoli"
        )
        params = []

        if self.current_filter:
            query += " WHERE argomento = ?"
            params.append(self.current_filter)
        elif self.current_search:
            query += " WHERE titolo_articolo LIKE ? OR sotto_titolo LIKE ?"
            params.extend([f"%{self.current_search}%", f"%{self.current_search}%"])

        query += " ORDER BY data DESC LIMIT ? OFFSET ?"
        params.extend([self.page_size, self.current_page * self.page_size])

        self.cursor.execute(query, params)
        return self.cursor.fetchall()

    def get_article_by_id(self, article_id):
        """Ritorna un articolo completo per ID"""
        self.cursor.execute(
            "SELECT * FROM t_articoli WHERE id_articolo = ?", (article_id,)
        )
        return self.cursor.fetchone()

    def get_argomenti(self):
        """Ritorna la lista degli argomenti con conteggio"""
        self.cursor.execute(
            """
            SELECT argomento, COUNT(*) as cnt
            FROM t_articoli
            GROUP BY argomento
            ORDER BY cnt DESC
        """
        )
        return self.cursor.fetchall()

    def clear_screen(self):
        """Pulisce lo schermo"""
        os.system("clear" if os.name == "posix" else "cls")

    def print_header(self):
        """Stampa l'intestazione"""
        print("\n" + "=" * 70)
        print("  📰 ESPLORA ARTICOLI - Database Ficiesse")
        print("=" * 70)

        if self.current_filter:
            print(f"  🏷️  Filtro attivo: {self.current_filter}")
        if self.current_search:
            print(f"  🔍 Ricerca: '{self.current_search}'")
        print()

    def print_articles_list(self, articles):
        """Stampa la lista degli articoli"""
        total = self.get_total_count()
        total_pages = (total + self.page_size - 1) // self.page_size

        page_info = (
            f"  Pagina {self.current_page + 1}/{total_pages}"
            f" ({total} articoli totali)\n"
        )
        print(page_info)
        print("-" * 70)

        for i, art in enumerate(articles, 1):
            data = art["data"][:10] if art["data"] else "N/D"
            argomento = (art["argomento"] or "N/D")[:14]
            titolo = (art["titolo_articolo"] or "N/D")[:20]
            exported = (
                art.get("esportato") if isinstance(art, dict) else art["esportato"]
            )
            status = "✅" if exported else " "

            left = (
                f"  {i:2}. [{art['id_articolo']}] {status} {data} | {argomento:<14} |"
            )
            print(left + f" {titolo}")

        print("-" * 70)

    def print_menu(self):
        """Stampa il menu"""
        print("\n  Comandi:")
        print("    [numero]  - Seleziona articolo (1-10)")
        print("    [n]ext    - Pagina successiva")
        print("    [p]rev    - Pagina precedente")
        print("    [f]iltra  - Filtra per argomento")
        print("    [s]earch  - Cerca per titolo")
        print("    [e]xport  - Esporta pagina corrente")
        print("    [a]ll     - Esporta risultati della ricerca (max 50)")
        print("    [r]eset   - Rimuovi filtri")
        print("    [q]uit    - Esci")
        print()

    def show_article_detail(self, article):
        """Mostra i dettagli di un articolo"""
        self.clear_screen()
        print("\n" + "=" * 70)
        print("  📄 DETTAGLIO ARTICOLO")
        print("=" * 70)

        print(f"\n  ID:          {article['id_articolo']}")
        print(f"  Data:        {article['data']}")
        print(f"  Argomento:   {article['argomento']}")
        print(f"  Titolo:      {article['titolo_articolo']}")
        print(f"  Sottotitolo: {article['sotto_titolo']}")
        print(f"  Visite:      {article['contatore_visite']}")

        # Anteprima testo (senza HTML)
        testo = article["testo_articolo"] or ""
        # Pulisci escape sequences
        testo = testo.replace("\\r\\n", "\n").replace("\\n", "\n").replace("\\r", "\n")
        testo = testo.replace("\r\n", "\n").replace("\r", "\n")

        if BS4_AVAILABLE:
            soup = BeautifulSoup(testo, "html.parser")
            testo_plain = soup.get_text(" ", strip=True)[:500]
        else:
            # Rimozione base dei tag HTML
            import re

            testo_plain = re.sub(r"<[^>]+>", " ", testo)
            testo_plain = re.sub(r"\s+", " ", testo_plain)[:500]

        print("\n  Anteprima testo:")
        print("-" * 70)
        # Dividi in righe da 68 caratteri
        for i in range(0, len(testo_plain), 68):
            print(f"  {testo_plain[i:i+68]}")
        if len(testo) > 500:
            print("  [...]")
        print("-" * 70)

        print("\n  Azioni:")
        print("    [e]xport  - Esporta in DOCX")
        print("    [b]ack    - Torna alla lista")
        print()

        while True:
            choice = input("  Scelta: ").strip().lower()

            if choice == "e":
                self.export_article(article)
            elif choice == "b":
                break
            else:
                print("  Comando non valido")

    def _clean_text_content(self, text):
        """Pulisce il testo convertendo escape sequences e normalizzando"""
        if not text:
            return ""

        # Converti sequenze di escape letterali in caratteri reali
        text = text.replace("\\r\\n", "\n")
        text = text.replace("\\n", "\n")
        text = text.replace("\\r", "\n")
        text = text.replace("\r\n", "\n")
        text = text.replace("\r", "\n")

        # Rimuovi righe vuote multiple
        import re

        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()

    def _clean_html_for_docx(self, html):
        """Pulisce l'HTML per renderlo compatibile con htmldocx"""
        if not html:
            return ""

        import re

        # Pulisci escape sequences
        html = self._clean_text_content(html)

        # Rimuovi tag non supportati da htmldocx
        unsupported_tags = [
            "INPUT",
            "TEXTAREA",
            "SELECT",
            "FORM",
            "SCRIPT",
            "STYLE",
            "IFRAME",
            "OBJECT",
            "EMBED",
        ]
        for tag in unsupported_tags:
            html = re.sub(
                rf"<{tag}[^>]*>.*?</{tag}>", "", html, flags=re.IGNORECASE | re.DOTALL
            )
            html = re.sub(rf"<{tag}[^>]*/?\s*>", "", html, flags=re.IGNORECASE)

        # Rimuovi attributi problematici (namespace XML, eventi JS)
        html = re.sub(r"<\?xml[^>]*\?>", "", html)
        html = re.sub(r'xmlns\s*=\s*["\'][^"\']*["\']', "", html)
        html = re.sub(r'on\w+\s*=\s*["\'][^"\']*["\']', "", html, flags=re.IGNORECASE)

        # Converti tag obsoleti
        html = re.sub(r"<FONT[^>]*>", "", html, flags=re.IGNORECASE)
        html = re.sub(r"</FONT>", "", html, flags=re.IGNORECASE)

        # Sistema tag vuoti che possono causare problemi
        html = re.sub(r"<([a-zA-Z]+)[^>]*>\s*</\1>", "", html)

        # Assicura che l'HTML sia wrappato in un tag root
        if not html.strip().startswith("<"):
            html = f"<p>{html}</p>"

        return html

    def _is_html_content(self, text):
        """Verifica se il testo contiene HTML significativo"""
        if not text:
            return False
        # Cerca tag HTML comuni (non solo <br> o &nbsp;)
        import re

        html_tags = re.findall(
            r"<(?:p|div|table|tr|td|h[1-6]|ul|ol|li|strong|em|b|i|a|span|font)[^>]*>",
            text,
            re.IGNORECASE,
        )
        return len(html_tags) > 0

    def export_article(self, article, interactive=True):
        """Esporta l'articolo in formato DOCX.

        Se interactive è False: non chiedere input; restituisci il percorso del file.
        """
        if not DOCX_AVAILABLE:
            if interactive:
                print("\n  ❌ Libreria 'python-docx' non installata!")
                print("     Installa con: pip install python-docx")
                input("\n  Premi INVIO per continuare...")
            else:
                logging.error("python-docx non installato; impossibile esportare")
            return None

        # Crea il documento
        doc = Document()

        # Titolo
        title = doc.add_heading(article["titolo_articolo"] or "Senza titolo", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sottotitolo
        if article["sotto_titolo"]:
            subtitle = doc.add_paragraph(article["sotto_titolo"])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in subtitle.runs:
                run.italic = True

        # Metadati
        doc.add_paragraph()
        meta = doc.add_paragraph()
        meta.add_run("Data: ").bold = True
        meta.add_run(str(article["data"] or "N/D"))
        meta.add_run("  |  ")
        meta.add_run("Argomento: ").bold = True
        meta.add_run(str(article["argomento"] or "N/D"))

        doc.add_paragraph()
        doc.add_paragraph("─" * 50)
        doc.add_paragraph()

        # Contenuto - pulisci prima il testo
        testo_raw = article["testo_articolo"] or ""
        testo_pulito = self._clean_text_content(testo_raw)

        # Verifica se è HTML vero o testo semplice
        if self._is_html_content(testo_pulito) and HTMLDOCX_AVAILABLE:
            try:
                testo_html = self._clean_html_for_docx(testo_pulito)
                parser = HtmlToDocx()
                parser.add_html_to_document(testo_html, doc)
            except Exception:
                self._add_plain_text(doc, testo_pulito)
        else:
            self._add_plain_text(doc, testo_pulito)

        # Salva il documento nella cartella export
        safe_title = "".join(
            c
            for c in (article["titolo_articolo"] or "articolo")
            if c.isalnum() or c in (" ", "-", "_")
        )[:50]
        filename = f"articolo_{article['id_articolo']}_{safe_title}.docx"

        # Crea la cartella export se non esiste
        export_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "export")
        os.makedirs(export_dir, exist_ok=True)
        filepath = os.path.join(export_dir, filename)

        try:
            doc.save(filepath)
            if interactive:
                print(f"\n  ✅ Articolo esportato: export/{filename}")
            else:
                logging.info(f"Esportato: export/{filename}")
        except Exception as e:
            if interactive:
                print(f"\n  ❌ Errore salvataggio: {e}")
            else:
                logging.error(f"Errore salvataggio: {e}")
            return None

        if interactive:
            input("\n  Premi INVIO per continuare...")

        return filepath

    def export_articles(self, articles, mark_exported=True, limit=50):
        """Esporta una lista di articoli (uno file DOCX per articolo).

        - Limita a `limit` elementi per chiamata per evitare blocchi di memoria.
        - Se mark_exported True, segna gli articoli come esportati nel DB.
        """
        total = len(articles)
        if total == 0:
            logging.info("Nessun articolo da esportare")
            return 0

        if total > limit:
            logging.warning(
                f"Tentativo di esportare {total} (limite {limit}); esporto {limit}."
            )
            articles = articles[:limit]
            total = limit

        # Use tqdm if available
        try:
            from tqdm import tqdm

            iterator = tqdm(articles, total=total, unit="articolo", ncols=80)
        except Exception:
            iterator = articles

        exported_count = 0
        for art in iterator:
            # art can be sqlite Row; convert to dict-like access
            article = art
            out = self.export_article(article, interactive=False)
            if out:
                exported_count += 1
                if mark_exported:
                    try:
                        self.cursor.execute(
                            "UPDATE t_articoli SET esportato = 1 WHERE id_articolo = ?",
                            (article["id_articolo"],),
                        )
                        self.conn.commit()
                    except Exception as e:
                        logging.error(
                            f"Errore nel marcare articolo {article['id_articolo']}: {e}"
                        )
        logging.info(f"Esportati: {exported_count}/{total} articoli")
        return exported_count

    def _add_plain_text(self, doc, text_content):
        """Aggiunge testo semplice al documento come paragrafi"""
        # Pulisci il contenuto
        text = self._clean_text_content(text_content)

        # Se contiene tag HTML, rimuovili
        if "<" in text:
            if BS4_AVAILABLE:
                soup = BeautifulSoup(text, "html.parser")
                text = soup.get_text("\n", strip=True)
            else:
                import re

                text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
                text = re.sub(r"</p>", "\n", text, flags=re.IGNORECASE)
                text = re.sub(r"<[^>]+>", "", text)

        # Decodifica entità HTML comuni
        text = text.replace("&nbsp;", " ")
        text = text.replace("&amp;", "&")
        text = text.replace("&lt;", "<")
        text = text.replace("&gt;", ">")
        text = text.replace("&quot;", '"')
        text = text.replace("&#39;", "'")

        # Pulisci spazi multipli
        import re

        text = re.sub(r" +", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Aggiungi ogni paragrafo
        for paragraph in text.split("\n"):
            paragraph = paragraph.strip()
            if paragraph:
                doc.add_paragraph(paragraph)

    def show_filter_menu(self):
        """Mostra il menu per filtrare per argomento"""
        self.clear_screen()
        print("\n" + "=" * 70)
        print("  🏷️  FILTRA PER ARGOMENTO")
        print("=" * 70 + "\n")

        argomenti = self.get_argomenti()

        for i, (arg, cnt) in enumerate(argomenti, 1):
            print(f"  {i:2}. {arg:<40} ({cnt} articoli)")

        print(f"\n  {0}. Rimuovi filtro")
        print()

        try:
            choice = input("  Seleziona numero: ").strip()
            if choice == "0":
                self.current_filter = None
                self.current_page = 0
            elif choice.isdigit():
                idx = int(choice) - 1
                if 0 <= idx < len(argomenti):
                    self.current_filter = argomenti[idx][0]
                    self.current_search = None
                    self.current_page = 0
        except (ValueError, IndexError):
            pass

    def run(self):
        """Avvia l'interfaccia interattiva"""
        if not self.connect():
            return

        try:
            while True:
                self.clear_screen()
                self.print_header()

                articles = self.get_articles_page()

                if not articles:
                    print("  Nessun articolo trovato.\n")
                else:
                    self.print_articles_list(articles)

                self.print_menu()

                choice = input("  Scelta: ").strip().lower()

                # Numero per selezionare articolo
                if choice.isdigit():
                    idx = int(choice) - 1
                    if 0 <= idx < len(articles):
                        article = self.get_article_by_id(articles[idx]["id_articolo"])
                        if article:
                            self.show_article_detail(article)

                # Navigazione
                elif choice in ("n", "next"):
                    total = self.get_total_count()
                    total_pages = (total + self.page_size - 1) // self.page_size
                    if self.current_page < total_pages - 1:
                        self.current_page += 1

                elif choice in ("p", "prev"):
                    if self.current_page > 0:
                        self.current_page -= 1

                # Filtri
                elif choice in ("f", "filtra"):
                    self.show_filter_menu()

                elif choice in ("s", "search"):
                    search = input("\n  Cerca nel titolo: ").strip()
                    if search:
                        self.current_search = search
                        self.current_filter = None
                        self.current_page = 0

                # Esporta pagina corrente
                elif choice in ("e", "export"):
                    if not articles:
                        print("\n  Nessun articolo da esportare in questa pagina")
                    else:
                        confirm = (
                            input(f"\n  Esportare {len(articles)} articoli? [y/N]: ")
                            .strip()
                            .lower()
                        )
                        if confirm == "y":
                            self.export_articles(
                                articles, mark_exported=True, limit=self.page_size
                            )

                # Esporta risultati (max 50)
                elif choice in ("a", "all"):
                    total_results = self.get_total_count()
                    if total_results == 0:
                        print("\n  Nessun articolo da esportare")
                    elif total_results > 50:
                        print(f"\n  Attenzione: {total_results} articoli (limite 50).")
                        confirm = (
                            input("  Vuoi esportare i primi 50 risultati? [y/N]: ")
                            .strip()
                            .lower()
                        )
                        if confirm == "y":
                            # fetch first 50 results
                            self.cursor.execute(
                                "SELECT * FROM t_articoli ORDER BY data DESC LIMIT 50"
                            )
                            rows = self.cursor.fetchall()
                            self.export_articles(rows, mark_exported=True, limit=50)
                    else:
                        self.cursor.execute(
                            "SELECT * FROM t_articoli ORDER BY data DESC"
                        )
                        rows = self.cursor.fetchall()
                        self.export_articles(rows, mark_exported=True, limit=50)

                elif choice in ("r", "reset"):
                    self.current_filter = None
                    self.current_search = None
                    self.current_page = 0

                # Esci
                elif choice in ("q", "quit", "exit"):
                    print("\n  Arrivederci! 👋\n")
                    break

        finally:
            self.close()


def check_dependencies():
    """Verifica e mostra lo stato delle dipendenze"""
    logging.info("\n📦 Verifica dipendenze:")
    logging.info(
        f"  - python-docx: {'✅ OK' if DOCX_AVAILABLE else '❌ Non installato'}"
    )
    hd_status = (
        "✅ OK"
        if HTMLDOCX_AVAILABLE
        else "⚠️ Non installato (export senza formattazione HTML)"
    )
    bs_status = "✅ OK" if BS4_AVAILABLE else "⚠️ Non installato (anteprima base)"
    logging.info("  - htmldocx:    " + hd_status)
    logging.info("  - beautifulsoup4: " + bs_status)

    if not DOCX_AVAILABLE:
        logging.warning("⚠️ Abilita export DOCX: installa python-docx, htmldocx")
        logging.info("   e installa beautifulsoup4 (bs4) per anteprime migliori")

    logging.info("")


def main():
    parser = argparse.ArgumentParser(description="Esplora il database degli articoli")
    parser.add_argument(
        "db", nargs="?", default="articoli.db", help="Percorso del database SQLite"
    )
    parser.add_argument(
        "--page-size", type=int, default=10, help="Numero di articoli per pagina"
    )
    parser.add_argument(
        "--verbose", "-v", action="store_true", help="Modalità verbosa (debug)"
    )
    parser.add_argument("--no-emoji", action="store_true", help="Output senza emoji")
    parser.add_argument(
        "--export-all",
        action="store_true",
        help="Esporta tutti gli articoli (non interattivo, limit applies)",
    )
    parser.add_argument(
        "--export-only-new",
        action="store_true",
        help="Esporta solo articoli non ancora esportati",
    )
    parser.add_argument(
        "--export-limit",
        type=int,
        default=50,
        help="Limite massimo articoli per export non-interattivo",
    )
    args = parser.parse_args()

    setup_logging(args.verbose, args.no_emoji)

    logging.info("\n" + "=" * 70)
    logging.info("  📰 ESPLORA ARTICOLI - Database Ficiesse")
    logging.info("=" * 70)

    check_dependencies()

    if not os.path.exists(args.db):
        logging.error(f"❌ Database '{args.db}' non trovato!")
        logging.info("   Usa: python esplora_articoli.py <database.db>")
        sys.exit(1)

    # Non-blocking: se export non interattivo, salta richiesta INVIO
    if not args.export_all:
        input("Premi INVIO per avviare l'esplorazione...")

    explorer = ArticoliExplorer(args.db)
    explorer.page_size = args.page_size

    # Non-interactive export flags
    if args.export_all:
        if not explorer.connect():
            sys.exit(1)
        try:
            only_new = args.export_only_new
            # fetch articles
            if only_new:
                explorer.cursor.execute(
                    (
                        "SELECT * FROM t_articoli WHERE esportato = 0 "
                        "ORDER BY data DESC LIMIT ?"
                    ),
                    (args.export_limit,),
                )
            else:
                explorer.cursor.execute(
                    "SELECT * FROM t_articoli ORDER BY data DESC LIMIT ?",
                    (args.export_limit,),
                )
            rows = explorer.cursor.fetchall()
            exported = explorer.export_articles(
                rows, mark_exported=True, limit=args.export_limit
            )
            logging.info(
                f"Export non-interattivo completato: {exported} articoli esportati"
            )
        finally:
            explorer.close()
        return

    explorer.run()


if __name__ == "__main__":
    main()
